import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Color palette ───────────────────────────────────────────────────────────
C_ACCENT  = RGBColor(0x00, 0x71, 0xE3)   # blue
C_DARK    = RGBColor(0x1D, 0x1D, 0x1F)   # near-black
C_GREY    = RGBColor(0x6E, 0x6E, 0x73)   # grey
C_GREEN   = RGBColor(0x18, 0x8A, 0x40)   # green
C_RED     = RGBColor(0xD9, 0x00, 0x00)   # red (danger questions)
C_ORANGE  = RGBColor(0xE5, 0x7A, 0x00)   # orange (warning)
C_BG_BLUE = RGBColor(0xEA, 0xF3, 0xFF)   # light blue bg
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Segoe UI"

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def para_font(para, size=11, bold=False, color=None, align=None):
    for run in para.runs:
        set_font(run, size=size, bold=bold, color=color)
    if align:
        para.alignment = align

def add_para(text="", size=11, bold=False, italic=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading(text, level=1):
    """Custom styled headings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=16, bold=True, color=C_ACCENT)
        # underline
        r.font.underline = False
    elif level == 2:
        set_font(r, size=13, bold=True, color=C_DARK)
    elif level == 3:
        set_font(r, size=11, bold=True, color=C_GREEN)
    return p

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0071E3')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, header_color="0071E3"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color=C_WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else "F5F5F7"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(cell_text))
            set_font(r, size=10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def qa_block(question, answer, danger=False):
    """Red question box + green answer."""
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run("  ВОПРОС: ")
    set_font(r, size=10, bold=True, color=C_WHITE)
    # Hack: use inline color via shading on paragraph — instead just bold colored text
    # Actually let's just do it simply
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.left_indent  = Cm(0.5)
    r1 = p2.add_run("Вопрос: ")
    set_font(r1, size=10.5, bold=True, color=C_RED if danger else C_ORANGE)
    r2 = p2.add_run(question)
    set_font(r2, size=10.5, bold=True, italic=True, color=C_DARK)

    # Answer
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(8)
    p3.paragraph_format.left_indent  = Cm(0.8)
    r3 = p3.add_run("Ответ: ")
    set_font(r3, size=10, bold=True, color=C_GREEN)
    r4 = p3.add_run(answer)
    set_font(r4, size=10, color=C_DARK)

def bullet(text, sub=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(1.2 if sub else 0.6)
    bullet_char = "–" if sub else "•"
    r1 = p.add_run(f"{bullet_char}  ")
    set_font(r1, size=10, bold=False, color=C_ACCENT)
    r2 = p.add_run(text)
    set_font(r2, size=10, color=C_DARK)

def term_block(term_en, term_ru, definition, in_project):
    """Styled block for each term."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{term_en}  ")
    set_font(r1, size=13, bold=True, color=C_ACCENT)
    r2 = p.add_run(f"({term_ru})")
    set_font(r2, size=11, bold=False, italic=True, color=C_GREY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.left_indent  = Cm(0.4)
    r3 = p2.add_run("Что это: ")
    set_font(r3, size=10, bold=True, color=C_DARK)
    r4 = p2.add_run(definition)
    set_font(r4, size=10, color=C_DARK)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after  = Pt(3)
    p3.paragraph_format.left_indent  = Cm(0.4)
    r5 = p3.add_run("В нашем проекте: ")
    set_font(r5, size=10, bold=True, color=C_GREEN)
    r6 = p3.add_run(in_project)
    set_font(r6, size=10, color=C_DARK)

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
add_para()
add_para()
p = add_para("ПОДГОТОВКА К ЗАЩИТЕ", size=22, bold=True, color=C_ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=6)
add_para("CS2 Match Outcome Prediction", size=14, italic=True, color=C_GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Narxoz University — Research Methods", size=11, color=C_GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Seisenbek Dias  ·  Mergen Temirzhan  ·  Onashov Aidos", size=10,
         color=C_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_divider()

add_para("Этот документ — твоя шпаргалка для защиты. Читай перед сном. "
         "Здесь все термины, объяснение моделей, расшифровка таблиц и самые "
         "опасные вопросы научрука с готовыми ответами.",
         size=10, italic=True, color=C_GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ЧАСТЬ 1 — КЛЮЧЕВЫЕ ТЕРМИНЫ
# ════════════════════════════════════════════════════════════════════════════
add_heading("ЧАСТЬ 1 — Ключевые термины", level=1)
add_para("Выучи каждый термин: что это, пример из жизни, и как это применимо к "
         "твоему проекту.", size=10, color=C_GREY, space_after=6)
add_divider()

# Dependent Variable
term_block(
    "Dependent Variable", "Зависимая переменная",
    "То, что ты предсказываешь. Результат. Зависит от других переменных.",
    "match_outcome — победил team1 или нет. 1 = победа, 0 = поражение."
)
qa_block(
    '"Почему именно это ваша зависимая переменная?"',
    "Потому что это конечный результат матча — именно его мы хотим предсказать. "
    "Всё остальное — статистика до матча — это то, что на него влияет."
)

add_divider()

# Independent Variables
term_block(
    "Independent Variables", "Независимые переменные",
    "То, что ты используешь для предсказания. Входные данные. "
    "Они не зависят от результата — они существуют до него.",
    "4 переменные: mean_hltv_rating, mean_kpr, head_to_head_win_rate, map_win_rate."
)

add_table(
    ["Переменная", "Что означает", "Диапазон"],
    [
        ["mean_hltv_rating",        "Средний рейтинг игроков команды (HLTV 2.0)",          "обычно 0.8 – 1.3"],
        ["mean_kpr",                "Среднее убийств за раунд (kills per round)",            "обычно 0.5 – 1.0"],
        ["head_to_head_win_rate",   "% предыдущих побед над этим конкретным оппонентом",    "0.0 – 1.0"],
        ["map_win_rate",            "% побед команды на конкретной карте матча",             "0.0 – 1.0"],
    ],
    col_widths=[5.0, 8.0, 3.5]
)

qa_block(
    '"Почему именно эти 4 переменные? Почему не больше?"',
    "Мы выбирали по трём критериям: доступность до начала матча, числовая измеримость, "
    "и теоретическое обоснование через гипотезы H1–H3. Добавлять переменные без "
    "теоретического обоснования — это data dredging, что снижает качество исследования."
)

add_divider()

# Research Question
term_block(
    "Research Question", "Исследовательский вопрос",
    "Главный вопрос, на который отвечает всё исследование. Должен быть конкретным и проверяемым.",
    '"Can pre-match team performance statistics predict the outcome of '
    'professional CS2 matches?" — Можно ли предсказать результат матча по '
    'статистике команды до игры?'
)

add_divider()

# Hypotheses
term_block(
    "Hypotheses", "Гипотезы",
    "Конкретные, проверяемые предположения. Должны быть falsifiable — "
    "то есть их можно опровергнуть данными.",
    "У нас 3 гипотезы: H1 (рейтинг/KPR), H2 (head-to-head), H3 (win rate на карте)."
)

add_table(
    ["Гипотеза", "Утверждение", "Подтвердилась?"],
    [
        ["H1", "Выше HLTV Rating / KPR → чаще побеждают",                "Частично — слабый эффект"],
        ["H2", "Лучший H2H рекорд → чаще побеждают в рематчах",          "Да — 2-й важный признак"],
        ["H3", "Лучший win rate на карте → чаще побеждают на ней",        "Да — самый важный признак"],
    ],
    col_widths=[2.5, 8.5, 4.0]
)

qa_block(
    '"Подтвердились ли ваши гипотезы?"',
    "Частично. Map win rate (H3) — самый важный предиктор, H3 подтверждена. "
    "Head-to-head (H2) — второй по важности, H2 подтверждена. HLTV Rating и KPR (H1) "
    "оказали меньший эффект, чем ожидалось. Модели достигли ~57% точности — это полезный, "
    "но неполный сигнал."
)

add_divider()

# Primary vs Secondary Data
add_heading("Primary Data  vs  Secondary Data", level=2)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r1 = p.add_run("Primary Data (Первичные данные): ")
set_font(r1, size=10, bold=True, color=C_DARK)
r2 = p.add_run("Данные, которые ты собрал сам — опросы, эксперименты, интервью.")
set_font(r2, size=10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(6)
r3 = p2.add_run("Secondary Data (Вторичные данные): ")
set_font(r3, size=10, bold=True, color=C_DARK)
r4 = p2.add_run("Данные, которые собрал кто-то другой, а ты их используешь.")
set_font(r4, size=10)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(3)
r5 = p3.add_run("В нашем проекте: ")
set_font(r5, size=10, bold=True, color=C_GREEN)
r6 = p3.add_run("Secondary data. Датасет с Kaggle (griffindesroches), "
                "собранный с HLTV.org. Мы не собирали данные сами.")
set_font(r6, size=10)

qa_block(
    '"Почему вторичные данные? Это слабость исследования?"',
    "Да, это ограничение — мы не контролируем процесс сбора. Однако собрать 7,033 "
    "матча самостоятельно было бы нереально в рамках курса, а HLTV.org — признанный "
    "профессиональный источник статистики CS2."
)

add_divider()

# Literature Review
term_block(
    "Literature Review", "Обзор литературы",
    "Раздел, где ты показываешь что уже существует по теме, "
    "и объясняешь как твоя работа на этом строится.",
    "Мы рассмотрели работы по CS:GO — Налапати (60–73% accuracy на in-game данных), "
    "Семёнофф (эконом-данные раундов). Наша работа первая на CS2 с pre-match данными."
)
qa_block(
    '"Что уже было сделано до вас в этой области?"',
    "Большинство работ изучали CS:GO, а не CS2. Налапати и др. (2018) достигли "
    "60–73% используя данные в реальном времени во время матча. Семёнофф — эконом-статистику "
    "раундов. Наша работа — первая с pre-match данными именно для CS2."
)

add_divider()

# Bias
add_heading("Bias (Смещение / Предвзятость)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Систематическая ошибка, которая искажает результаты в одну сторону. "
              "Есть в каждом исследовании — важно их признавать, а не скрывать.")
set_font(r, size=10, color=C_DARK)

add_table(
    ["Тип Bias", "Что происходит", "Почему это проблема"],
    [
        ["Selection Bias",
         "Только профессиональные матчи",
         "Результаты не переносятся на любительский уровень"],
        ["Perspective Bias",
         "Все записи с точки зрения team1",
         "Создаёт асимметрию в данных"],
        ["Temporal Bias",
         "Данные 2024–2025, патчи меняют игру",
         "Старая статистика может быть нерелевантна"],
        ["Survivorship Bias",
         "Только команды, игравшие на этом уровне",
         "Слабые команды, выбывшие раньше, не представлены"],
    ],
    col_widths=[4.0, 6.0, 6.0]
)

qa_block(
    '"Какие bias присутствуют в вашей работе?"',
    "Selection bias (только проф. матчи), perspective bias (данные только team1), "
    "temporal bias (патчи CS2 меняют мету), survivorship bias (только активные команды). "
    "Мы признаём эти ограничения в разделе Limitations.",
    danger=True
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ЧАСТЬ 2 — КАК ОБУЧАЛИСЬ МОДЕЛИ
# ════════════════════════════════════════════════════════════════════════════
add_heading("ЧАСТЬ 2 — Как обучались модели (объяснение простыми словами)", level=1)
add_divider()

add_heading("Шаг 1 — Разделение данных", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Из 7,033 матчей мы разделили:")
set_font(r, size=10)
bullet("80% = 5,626 матчей → Train Set (обучение) — модели видят эти данные")
bullet("20% = 1,407 матчей → Test Set (проверка) — модели никогда не видели эти данные")

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(6)
p2.paragraph_format.space_before = Pt(4)
r2 = p2.add_run("Аналогия: как подготовка к экзамену (train) и сам экзамен (test). "
                "Нечестно давать на экзамене те же задачи что были на подготовке.")
set_font(r2, size=10, italic=True, color=C_GREY)

add_heading("Шаг 2 — Четыре модели", level=2)

# Naive Baseline
add_heading("Naive Baseline — Тупая базовая линия", level=3)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Не обучается вообще. Просто всегда говорит «team1 победит» — потому что "
              "в 54.6% случаев team1 и правда побеждает. Это наш минимальный порог: "
              "любая нормальная модель должна быть лучше этого.")
set_font(r, size=10)

# Logistic Regression
add_heading("Logistic Regression — Логистическая регрессия", level=3)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Рисует одну прямую линию (в пространстве 4 переменных), которая разделяет "
              "«победы» и «поражения». Простая и интерпретируемая, но не ловит сложные паттерны.")
set_font(r, size=10)

# Random Forest
add_heading("Random Forest — Случайный лес  ★ Победитель", level=3)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Строит много (100+) маленьких деревьев решений. Каждое дерево как мини-эксперт: "
              "«если map_win_rate > 0.6 и h2h > 0.5 → победа». Каждое обучается на "
              "случайном подмножестве данных. Итоговый ответ — голосование всех деревьев. "
              "Устойчив к переобучению.")
set_font(r, size=10)

# Gradient Boosting
add_heading("Gradient Boosting — Градиентный бустинг", level=3)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Тоже деревья, но умнее. Строит последовательно: первое дерево делает предсказание, "
              "второе учится на ошибках первого, третье — на ошибках второго. "
              "Каждый следующий исправляет предыдущего.")
set_font(r, size=10)

qa_block(
    '"Почему Random Forest победил?"',
    "Random Forest лучше справляется с нелинейными зависимостями в данных без переобучения, "
    "благодаря усреднению по многим деревьям. Однако разница между моделями статистически "
    "незначима (p > 0.05 на paired t-test), так что победа символическая.",
    danger=False
)

qa_block(
    '"Что такое overfitting и есть ли он у вас?"',
    "Overfitting — когда модель запоминает обучающие данные вместо того чтобы учить паттерны. "
    "У нас нет значительного overfitting: мы проверяли на отдельном test set (20% данных) "
    "и через 5-fold cross-validation.",
    danger=True
)

qa_block(
    '"Что такое 5-fold cross-validation?"',
    "Делим данные на 5 равных частей. Обучаем модель 5 раз: каждый раз 4 части — обучение, "
    "1 часть — тест. Итог — среднее по 5 запускам. Это надёжнее одного split, "
    "потому что показывает стабильность модели.",
    danger=False
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ЧАСТЬ 3 — ТАБЛИЦА РЕЗУЛЬТАТОВ
# ════════════════════════════════════════════════════════════════════════════
add_heading("ЧАСТЬ 3 — Расшифровка таблицы результатов", level=1)
add_divider()

add_heading("Что означает каждый столбец", level=2)

add_table(
    ["Столбец", "Что измеряет", "Простыми словами", "Идеальное значение"],
    [
        ["Accuracy",
         "% правильных предсказаний",
         "Из 100 матчей сколько угадал",
         "100%"],
        ["Precision",
         "Из предсказанных побед — сколько реальных побед",
         "Насколько ты уверен когда говоришь 'победят'",
         "100%"],
        ["Recall",
         "Из реальных побед — сколько ты их поймал",
         "Насколько полно ты находишь победы",
         "100%"],
        ["Macro F1",
         "Среднее Precision и Recall для обоих классов",
         "Общая сбалансированность модели",
         "1.00"],
        ["AUC-ROC",
         "Способность различать победу от поражения",
         "0.5 = случайно, 1.0 = идеально",
         "1.00"],
        ["5-Fold CV Acc",
         "Accuracy на 5 разных разбивках данных",
         "Насколько стабильна модель",
         "100%"],
    ],
    col_widths=[3.2, 4.5, 5.0, 3.0]
)

add_heading("Итоговые результаты наших моделей", level=2)

add_table(
    ["Модель", "Accuracy", "Precision", "Recall", "Macro F1", "AUC-ROC"],
    [
        ["Naive Baseline",      "54.6%", "—",     "—",     "0.353", "0.500"],
        ["Logistic Regression", "56.2%", "0.563", "0.562", "0.484", "0.574"],
        ["Random Forest ★",     "57.2%", "0.570", "0.572", "0.536", "0.571"],
        ["Gradient Boosting",   "56.4%", "0.564", "0.564", "0.528", "0.573"],
    ],
    col_widths=[4.5, 2.2, 2.2, 2.2, 2.2, 2.2]
)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("★ Random Forest — победитель по Macro F1 (0.536). "
              "Разница между моделями статистически незначима (p > 0.05).")
set_font(r, size=10, bold=True, color=C_GREEN)

qa_block(
    '"57% — это хороший результат? Почему так мало?"',
    "В спортивном прогнозировании даже 55–60% считается хорошим результатом, потому что "
    "результат зависит от множества факторов в реальном времени: психологическое состояние, "
    "стратегии, технические проблемы. Наши 4 переменные — только часть картины. "
    "Принципиально используем только pre-match данные, что ограничивает потолок точности.",
    danger=True
)

qa_block(
    '"Почему Macro F1, а не просто Accuracy?"',
    "Потому что у нас небалансированные классы: 54.6% побед против 45.4% поражений. "
    "Модель, всегда говорящая 'победа', получила бы 54.6% Accuracy, но была бы бесполезной. "
    "Macro F1 учитывает оба класса одинаково.",
    danger=True
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ЧАСТЬ 4 — САМЫЕ ОПАСНЫЕ ВОПРОСЫ
# ════════════════════════════════════════════════════════════════════════════
add_heading("ЧАСТЬ 4 — Самые опасные вопросы (читай особенно внимательно)", level=1)
add_divider()

add_para("Эти вопросы — любимые у научруков. Если ты ответишь на них уверенно, "
         "оценка будет высокой.", size=10, italic=True, color=C_GREY, space_after=6)

qa_block(
    '"Ваша выборка репрезентативна?"',
    "Для профессиональных CS2 матчей с мая 2024 — да. 7,033 матчей из 648 турниров "
    "охватывают большинство крупных соревнований. Для любительских матчей — нет, "
    "и мы это признаём как ограничение в разделе Limitations.",
    danger=True
)

qa_block(
    '"Почему не нейросети? Почему не Deep Learning?"',
    "Нейросети требуют значительно больше данных и сложнее интерпретировать. "
    "С 7,033 записями и 4 переменными классические методы более подходящие. "
    "Интерпретируемость — ключевое требование в академическом исследовании.",
    danger=True
)

qa_block(
    '"В чём практическая ценность 57% точности?"',
    "Прямая коммерческая ценность минимальна — именно поэтому мы явно указываем: "
    "не для беттинга, не для рекрутинга. Ценность академическая: мы показали, "
    "что map win rate и head-to-head содержат предиктивный сигнал, что направляет "
    "будущие исследования с более богатыми данными.",
    danger=True
)

qa_block(
    '"Что такое paired t-test и зачем вы его использовали?"',
    "Парный t-тест проверяет, является ли разница между двумя наборами чисел "
    "статистически значимой или это случайность. Мы сравнивали accuracy моделей "
    "на 5 fold'ах. Результат p > 0.05 означает: разница между моделями не значима, "
    "то есть Random Forest победил 'на удачу', а не системно.",
    danger=False
)

qa_block(
    '"Почему mean imputation для пропущенных значений?"',
    "63 матча не имели данных о map_win_rate — команда не играла на этой карте раньше. "
    "Mean imputation (заменить средним значением) — стандартный метод, сохраняющий "
    "распределение без удаления строк. Удаление 63 строк не изменило бы результат "
    "существенно, но ввело бы selection bias.",
    danger=False
)

qa_block(
    '"Что такое Winsorization и почему вы её применили?"',
    "Winsorization — это обрезка выбросов: значения ниже 1-го перцентиля заменяются "
    "1-м перцентилем, выше 99-го — 99-м. Мы применили для всех 4 переменных, "
    "чтобы экстремальные значения (например, команда с 100% win rate на карте из "
    "2 матчей) не искажали модель.",
    danger=False
)

qa_block(
    '"Что такое StandardScaler и зачем он нужен?"',
    "StandardScaler нормализует переменные — приводит их к среднему 0 и "
    "стандартному отклонению 1. Это нужно потому что переменные в разных масштабах: "
    "KPR около 0.7, а map_win_rate от 0 до 1. Без нормализации модель думает, "
    "что большие числа важнее. Важно: fitted только на train set, чтобы избежать data leakage.",
    danger=False
)

qa_block(
    '"Почему Random Forest — лучшая модель, а не Gradient Boosting?"',
    "Технически Gradient Boosting часто мощнее Random Forest. "
    "Но у нас разница незначима статистически (p > 0.05). "
    "Random Forest выиграл по Macro F1 (0.536 vs 0.528) — это наша ключевая метрика. "
    "При несущественной разнице выбор можно обосновать любой из двух.",
    danger=False
)

# ════════════════════════════════════════════════════════════════════════════
# ЧАСТЬ 5 — БЫСТРАЯ ШПАРГАЛКА
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("ЧАСТЬ 5 — Быстрая шпаргалка (выучи наизусть)", level=1)
add_divider()

add_heading("Ключевые числа проекта", level=2)
add_table(
    ["Параметр", "Значение"],
    [
        ["Количество матчей",               "7,033"],
        ["Количество турниров",              "648"],
        ["Период данных",                    "Май 2024 – Октябрь 2025"],
        ["Источник данных",                  "HLTV.org через Kaggle (griffindesroches)"],
        ["Train / Test split",               "80% / 20%  (random_state=42, stratified)"],
        ["Пропуски в map_win_rate",           "63 матча → mean imputation"],
        ["Баланс классов",                   "54.6% победа / 45.4% поражение"],
        ["Лучшая модель",                    "Random Forest"],
        ["Лучший Macro F1",                  "0.536"],
        ["Лучший Accuracy",                  "57.2%"],
        ["Лучший AUC-ROC",                   "0.571"],
        ["Самый важный признак",             "map_win_rate"],
        ["2-й по важности признак",          "head_to_head_win_rate"],
        ["Статистическая значимость",        "p > 0.05 → разница между моделями незначима"],
    ],
    col_widths=[7.0, 9.0]
)

add_heading("Одна фраза на каждый термин", level=2)
add_table(
    ["Термин", "Одна фраза"],
    [
        ["Dependent Variable",       "То, что предсказываем = match_outcome (1/0)"],
        ["Independent Variable",     "То, чем предсказываем = 4 признака команды"],
        ["Research Question",        "Можно ли по статистике до матча предсказать победителя?"],
        ["Hypothesis",               "Проверяемое предположение — H1, H2, H3"],
        ["Primary Data",             "Собрал сам — опрос, эксперимент"],
        ["Secondary Data",           "Взял у других — у нас это HLTV через Kaggle"],
        ["Literature Review",        "Обзор что уже сделано — CS:GO работы, Налапати, Семёнофф"],
        ["Bias",                     "Систематическая ошибка — selection, temporal, perspective"],
        ["Overfitting",              "Модель запомнила данные вместо паттернов — мы избежали"],
        ["Cross-validation",         "Проверяем модель 5 раз на разных частях данных"],
        ["Macro F1",                 "Главная метрика — баланс precision и recall для обоих классов"],
        ["AUC-ROC",                  "Насколько хорошо модель разделяет классы (0.5 = случайно)"],
        ["Winsorization",            "Обрезка выбросов — заменяем экстремальные значения"],
        ["StandardScaler",           "Нормализация — приводим все переменные к одному масштабу"],
    ],
    col_widths=[4.5, 11.0]
)

# ── Save ────────────────────────────────────────────────────────────────────
out_path = r"final\defense_preparation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
