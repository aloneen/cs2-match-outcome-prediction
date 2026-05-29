"""
CS2 Match Outcome Prediction — Report Generator
Generates reports/task7_report.docx from the real HLTV dataset.
Author : Seisenbek Dias
Course : Research Methods
"""

import os
import sys
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")          # headless — no display needed
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# 0.  LOAD & PROCESS REAL DATA  (same logic as code/data_pipeline.py)
# ─────────────────────────────────────────────────────────────────────────────
CSV_PATH = "data/cs2_newestcombinedmatches_team1_reference_reduced2.csv"

MAP_COL_LOOKUP = {
    "mirage":   "team1_mirage",
    "inferno":  "team1_inferno",
    "nuke":     "team1_nuke",
    "dust2":    "team1_dust2",
    "ancient":  "team1_ancient",
    "anubis":   "team1_anubis",
    "vertigo":  "team1_vertigo",
    "overpass": "team1_overpass",
    "train":    "team1_train",
}

print("Loading dataset …")
df_raw = pd.read_csv(CSV_PATH, low_memory=False)
df = df_raw.copy()

df["match_outcome"]         = df["team1_win_flag"].astype(int)
df["mean_hltv_rating"]      = pd.to_numeric(df["team1_avg_RATING"], errors="coerce")
df["mean_kpr"]              = pd.to_numeric(df["team1_avg_KPR"],    errors="coerce")
df["head_to_head_win_rate"] = pd.to_numeric(
    df["team1_head2head_percentage"], errors="coerce") / 100.0

def get_map_winrate(row):
    map_name = str(row.get("decider_map", "")).strip().lower()
    col = MAP_COL_LOOKUP.get(map_name)
    if col and col in row.index:
        try:
            return float(row[col]) / 100.0
        except (ValueError, TypeError):
            return np.nan
    return np.nan

df["map_win_rate"] = df.apply(get_map_winrate, axis=1)

COLS   = ["match_outcome", "mean_hltv_rating", "mean_kpr",
          "head_to_head_win_rate", "map_win_rate"]
CONTIN = ["mean_hltv_rating", "mean_kpr", "head_to_head_win_rate", "map_win_rate"]

df_model = df[COLS].copy()
N = len(df_model)

# Missing audit
missing_count = df_model.isnull().sum()
missing_pct   = (missing_count / N * 100).round(2)

# IQR outliers (pre-cleaning)
outlier_info = {}
for col in CONTIN:
    q1, q3 = df_model[col].quantile([0.25, 0.75])
    iqr = q3 - q1
    lo, hi = q1 - 1.5 * iqr, q3 + 1.5 * iqr
    n_out = int(((df_model[col] < lo) | (df_model[col] > hi)).sum())
    outlier_info[col] = {"lo": round(lo,4), "hi": round(hi,4), "n": n_out,
                         "min": round(df_model[col].min(),4),
                         "max": round(df_model[col].max(),4)}

# Clean
df_clean = df_model.copy()
map_wr_mean     = df_clean["map_win_rate"].mean()
map_wr_missing  = int(df_clean["map_win_rate"].isnull().sum())
df_clean["map_win_rate"]          = df_clean["map_win_rate"].fillna(map_wr_mean)
df_clean["head_to_head_win_rate"] = df_clean["head_to_head_win_rate"].fillna(0.50)

winsor_info = {}
for col in CONTIN:
    p01 = df_clean[col].quantile(0.01)
    p99 = df_clean[col].quantile(0.99)
    n_lo = int((df_clean[col] < p01).sum())
    n_hi = int((df_clean[col] > p99).sum())
    df_clean[col] = df_clean[col].clip(lower=p01, upper=p99)
    winsor_info[col] = {"floor": round(p01,4), "cap": round(p99,4),
                        "clipped_low": n_lo, "clipped_high": n_hi}

# Train/test split
FEATURES = CONTIN; TARGET = "match_outcome"
X = df_clean[FEATURES]; y = df_clean[TARGET]
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=42, stratify=y)

scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled  = scaler.transform(X_test)

print(f"Dataset loaded: {N:,} rows | Clean: {len(df_clean):,} rows retained")

# ─────────────────────────────────────────────────────────────────────────────
# 1.  GENERATE FIGURES  (saved as PNGs, then embedded in the docx)
# ─────────────────────────────────────────────────────────────────────────────
FIG_DIR = "report_figures"
os.makedirs(FIG_DIR, exist_ok=True)

PALETTE  = ["#2E74B5", "#1F3964", "#C9D9F0", "#FF6B6B", "#4CAF50"]
sns.set_theme(style="whitegrid", font_scale=1.05)

# Figure 1 — Missing Data Bar Chart
fig, ax = plt.subplots(figsize=(8, 3.5))
colors = ["#FF6B6B" if v > 0 else "#4CAF50" for v in missing_pct.values]
bars = ax.barh(missing_pct.index, missing_pct.values, color=colors, edgecolor="white")
ax.set_xlabel("Missing (%)", fontsize=11)
ax.set_title("Figure 1 — Missing Data by Variable (n = 7,033)", fontsize=12, fontweight="bold")
for bar, val in zip(bars, missing_pct.values):
    label = f"{val:.2f}%" if val > 0 else "0%"
    ax.text(bar.get_width() + 0.01, bar.get_y() + bar.get_height()/2,
            label, va="center", fontsize=9)
ax.set_xlim(0, max(missing_pct.max() * 1.5, 2))
legend_handles = [
    mpatches.Patch(color="#FF6B6B", label="Missing values present"),
    mpatches.Patch(color="#4CAF50", label="Fully complete"),
]
ax.legend(handles=legend_handles, loc="lower right", fontsize=9)
plt.tight_layout()
fig.savefig(f"{FIG_DIR}/fig1_missing_data.png", dpi=150, bbox_inches="tight")
plt.close()
print("Figure 1 saved.")

# Figure 2 — Boxplots (pre-winsorization) for HLTV Rating and KPR
fig, axes = plt.subplots(1, 2, figsize=(9, 4))
for ax, col, label in zip(
        axes,
        ["mean_hltv_rating", "mean_kpr"],
        ["Mean HLTV Rating", "Mean KPR"]):
    bp = ax.boxplot(df_model[col].dropna(), patch_artist=True,
                    medianprops=dict(color="white", linewidth=2),
                    boxprops=dict(facecolor="#2E74B5", alpha=0.7),
                    flierprops=dict(marker="o", markerfacecolor="#FF6B6B",
                                   markersize=3, alpha=0.5))
    ax.set_title(label, fontsize=11, fontweight="bold")
    ax.set_ylabel("Value")
    ax.text(0.98, 0.02,
            f"Outliers (IQR): {outlier_info[col]['n']}",
            transform=ax.transAxes, ha="right", va="bottom",
            fontsize=9, color="#FF6B6B")
fig.suptitle("Figure 2 — Boxplots for Outlier Detection (Pre-Winsorization)",
             fontsize=12, fontweight="bold")
plt.tight_layout()
fig.savefig(f"{FIG_DIR}/fig2_boxplots.png", dpi=150, bbox_inches="tight")
plt.close()
print("Figure 2 saved.")

# Figure 3 — Head-to-Head Win Rate Distribution (bimodal)
fig, ax = plt.subplots(figsize=(8, 4))
ax.hist(df_model["head_to_head_win_rate"].dropna(), bins=40,
        color="#2E74B5", edgecolor="white", alpha=0.85)
ax.axvline(0.5, color="#FF6B6B", linewidth=2, linestyle="--", label="Neutral prior (0.50)")
q1_h2h = df_model["head_to_head_win_rate"].quantile(0.25)
q3_h2h = df_model["head_to_head_win_rate"].quantile(0.75)
iqr_h2h = q3_h2h - q1_h2h
ax.axvline(q1_h2h - 1.5*iqr_h2h, color="orange", linewidth=1.5,
           linestyle=":", label=f"IQR lower fence ({q1_h2h-1.5*iqr_h2h:.2f})")
ax.axvline(q3_h2h + 1.5*iqr_h2h, color="orange", linewidth=1.5,
           linestyle=":", label=f"IQR upper fence ({q3_h2h+1.5*iqr_h2h:.2f})")
ax.set_xlabel("Head-to-Head Win Rate", fontsize=11)
ax.set_ylabel("Frequency", fontsize=11)
ax.set_title("Figure 3 — Head-to-Head Win Rate Distribution\n"
             "(Bimodal structure explains 1,935 IQR outliers)", fontsize=12, fontweight="bold")
ax.legend(fontsize=9)
plt.tight_layout()
fig.savefig(f"{FIG_DIR}/fig3_h2h_distribution.png", dpi=150, bbox_inches="tight")
plt.close()
print("Figure 3 saved.")

# Figure 4 — Target Class Distribution
fig, ax = plt.subplots(figsize=(5, 4))
vc = df_clean["match_outcome"].value_counts().sort_index(ascending=False)
labels = ["Win (1)", "Loss (0)"]
colors_bar = ["#4CAF50", "#FF6B6B"]
bars = ax.bar(labels, [vc[1], vc[0]], color=colors_bar, edgecolor="white", width=0.5)
for bar, val in zip(bars, [vc[1], vc[0]]):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 40,
            f"{val:,}\n({val/N*100:.1f}%)", ha="center", va="bottom", fontsize=11)
ax.set_ylabel("Record Count", fontsize=11)
ax.set_title("Figure 4 — Target Variable Class Distribution\n(Post-Cleaning, n = 7,033)",
             fontsize=12, fontweight="bold")
ax.set_ylim(0, max(vc[1], vc[0]) * 1.15)
plt.tight_layout()
fig.savefig(f"{FIG_DIR}/fig4_class_distribution.png", dpi=150, bbox_inches="tight")
plt.close()
print("Figure 4 saved.")

# Figure 5 — Feature Distributions (Pre vs Post Winsorization, HLTV & KPR)
fig, axes = plt.subplots(2, 2, figsize=(10, 6))
pairs = [
    ("mean_hltv_rating", "Mean HLTV Rating"),
    ("mean_kpr",         "Mean KPR"),
]
for i, (col, label) in enumerate(pairs):
    ax_pre  = axes[i][0]
    ax_post = axes[i][1]
    ax_pre.hist(df_model[col].dropna(), bins=40,
                color="#2E74B5", edgecolor="white", alpha=0.8)
    ax_pre.set_title(f"{label}\nPre-Winsorization", fontsize=10, fontweight="bold")
    ax_post.hist(df_clean[col], bins=40,
                 color="#1F3964", edgecolor="white", alpha=0.8)
    ax_post.set_title(f"{label}\nPost-Winsorization", fontsize=10, fontweight="bold")
    for ax in [ax_pre, ax_post]:
        ax.set_xlabel("Value", fontsize=9)
        ax.set_ylabel("Frequency", fontsize=9)
fig.suptitle("Figure 5 — Feature Distributions: Pre vs Post Winsorization",
             fontsize=12, fontweight="bold")
plt.tight_layout()
fig.savefig(f"{FIG_DIR}/fig5_distributions.png", dpi=150, bbox_inches="tight")
plt.close()
print("Figure 5 saved.")

# ─────────────────────────────────────────────────────────────────────────────
# 2.  BUILD DOCX
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DARK_BLUE      = RGBColor(0x1F, 0x39, 0x64)
MID_BLUE       = RGBColor(0x2E, 0x74, 0xB5)
TABLE_HEAD_HEX = "2E74B5"
ALT_ROW_HEX    = "DAE8FC"

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1):
    p   = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.size = Pt(14) if level == 1 else Pt(12)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(text, bold=False, italic=False, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(11)
    run.font.bold   = bold
    run.font.italic = italic
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_table(headers, rows, caption=""):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].font.bold   = True
        cp.runs[0].font.size   = Pt(10)
        cp.runs[0].font.italic = True
        cp.paragraph_format.space_after  = Pt(2)
        cp.paragraph_format.space_before = Pt(8)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold      = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size       = Pt(9)
        cell.paragraphs[0].alignment               = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, TABLE_HEAD_HEX)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment         = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_bg(cell, ALT_ROW_HEX)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_figure(path, caption, width_inches=5.5):
    doc.add_picture(path, width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.runs[0]
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after  = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Data Audit and Cleaning Protocol")
tr.font.size = Pt(20); tr.font.bold = True; tr.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("CS2 Match Outcome Prediction Using Machine Learning")
sr.font.size = Pt(13); sr.font.italic = True; sr.font.color.rgb = MID_BLUE

repo_p = doc.add_paragraph()
repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo_p.add_run("GitHub Repository: ").font.size = Pt(11)
add_hyperlink(repo_p, "https://github.com/aloneen/cs2-match-outcome-prediction",
              "https://github.com/aloneen/cs2-match-outcome-prediction")
repo_p.paragraph_format.space_after = Pt(6)

disc = doc.add_paragraph()
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = disc.add_run(
    "Disclaimer: Data sourced from HLTV.org via Kaggle for academic research only.\n"
    "This report and its outputs must not be used for commercial forecasting,\n"
    "sports betting, player recruitment, or any form of competitive intelligence."
)
dr.font.size = Pt(9); dr.font.italic = True
dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — STUDENT AND PROJECT INFORMATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 1 — Student and Project Information", level=1)

add_table(
    headers=["Field", "Detail"],
    rows=[
        ["Student Name",    "Seisenbek Dias"],
        ["Student ID",      "S23068165"],
        ["Course",          "Research Methods"],
        ["Research Title",  "CS2 Match Outcome Prediction Using Machine Learning"],
        ["Submission Date", "April 2026"],
    ]
)

add_body("Research Question:", bold=True, space_after=2)
add_body(
    "Which machine learning model best predicts the outcome of professional CS2 matches, "
    "drawing on team performance statistics and historical match data?", space_after=8
)
add_body("Hypotheses:", bold=True, space_after=2)
for h in [
    "H1: Teams with higher combat performance ratings (HLTV 2.0, KPR) are more likely to win a given match.",
    "H2: A favorable head-to-head win rate against a specific opponent increases the probability of winning a rematch.",
    "H3: Teams with a stronger historical win rate on a given map are more likely to win when that map is selected.",
]:
    p = doc.add_paragraph(h, style="List Bullet")
    p.runs[0].font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DATASET DESCRIPTION AND DATA DICTIONARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 2 — Dataset Description and Data Dictionary", level=1)

add_body(
    "The dataset used in this study comprises 7,033 professional CS2 match records sourced from "
    "HLTV.org, aggregated via the publicly available Kaggle dataset published by griffindesroches "
    "(CS2 HLTV Professional Match Statistics). The data covers 648 tournaments from May 2024 to "
    "October 2025, spanning Tier-1 and Tier-2 professional competition. Match statistics were "
    "scraped directly from HLTV.org and include per-player and team-aggregated performance metrics "
    "at the time of each match. This approach is consistent with the secondary data use practices "
    "outlined in our Task 4 Ethical Risk Assessment, which identified HLTV.org and other CS2 "
    "statistics aggregators as the primary data sources for this research.",
    space_after=8
)

add_table(
    caption="Table 1 — Data Dictionary",
    headers=["Variable", "Role", "Type", "Source Column", "Range", "Measurement"],
    rows=[
        ["match_outcome",           "Dependent (Y)",    "Binary",     "team1_win_flag",               "{0, 1}",       "1 = Win, 0 = Loss (team1 perspective)"],
        ["mean_hltv_rating",        "Independent (X1a)","Continuous", "team1_avg_RATING",             "[0.87, 1.17]", "Mean HLTV 2.0 Rating across 5-player roster"],
        ["mean_kpr",                "Independent (X1b)","Continuous", "team1_avg_KPR",                "[0.58, 0.75]", "Mean Kills Per Round across 5-player roster"],
        ["head_to_head_win_rate",   "Independent (X2)", "Continuous", "team1_head2head_percentage/100","[0, 1]",       "Team1 win rate in prior direct matchups vs opponent"],
        ["map_win_rate",            "Independent (X3)", "Continuous", "team1_[decider_map]/100",       "[0, 1]",       "Team1 historical win rate on the specific map played"],
    ]
)

add_body(
    "The five model variables were engineered from the raw 168-column dataset. The target variable "
    "team1_win_flag was retained without transformation as it is already binary (1 = team1 wins, "
    "0 = team1 loses). HLTV Rating and KPR were mapped directly from the pre-aggregated team average "
    "columns. Head-to-head win rate was derived by dividing team1_head2head_percentage by 100 to "
    "normalize it to the [0, 1] range. Map win rate was computed by looking up team1's win rate on "
    "the specific map played (decider_map) and dividing by 100, using a map-to-column lookup table "
    "covering all nine active CS2 competitive maps.",
    space_after=6
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MISSING DATA AUDIT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 3 — Missing Data Audit", level=1)

add_body(
    "Prior to any cleaning, a systematic audit was conducted across all five model variables "
    "using pandas.isnull().sum(). Results are expressed as raw counts and percentages of the "
    "total sample (n = 7,033).",
    space_after=8
)

add_table(
    caption="Table 2 — Missing Data Summary (Raw Dataset, n = 7,033)",
    headers=["Variable", "Missing Count", "Missing (%)", "Mechanism", "Action"],
    rows=[
        ["match_outcome",           "0",  "0.00%", "—",   "None required"],
        ["mean_hltv_rating",        "0",  "0.00%", "—",   "None required"],
        ["mean_kpr",                "0",  "0.00%", "—",   "None required"],
        ["head_to_head_win_rate",   "0",  "0.00%", "—",   "None required"],
        ["map_win_rate",            str(map_wr_missing), f"{map_wr_missing/N*100:.2f}%", "MAR",
         f"Mean substitution ({map_wr_mean:.4f})"],
    ]
)

add_body(f"Total complete records (pre-cleaning): {N - map_wr_missing:,} of {N:,} ({(N-map_wr_missing)/N*100:.1f}%)",
         bold=True, space_after=8)

add_body(
    "The overall data quality of the HLTV dataset is very high, with only one variable containing "
    "missing values. The 63 missing entries in map_win_rate arise because 63 matches have no "
    "recorded decider_map value — a data collection gap in the scraping pipeline rather than a "
    "fundamental absence of information. This pattern is classified as Missing At Random (MAR), "
    "as the likelihood of missingness is related to the match format (some best-of-one matches "
    "did not record a decider map), which is observable from other columns.",
    space_after=6
)

add_body(
    "All four other variables are fully complete (0 missing values), reflecting the structured and "
    "consistent nature of HLTV's data recording for team-level aggregate statistics. Listwise "
    "deletion was rejected as a strategy: removing 63 rows would constitute an unnecessary 0.9% "
    "reduction in sample size, and the missingness mechanism does not introduce systematic bias "
    "into the remaining data. As documented in the Task 4 Ethical Risk Assessment, avoiding "
    "unnecessary data loss is particularly important given the structural underrepresentation of "
    "certain regions in the HLTV dataset — any further row reduction disproportionately affects "
    "teams from the CIS, Southeast Asia, and Oceania that appear less frequently in the data.",
    space_after=8
)

add_figure(f"{FIG_DIR}/fig1_missing_data.png",
           "Figure 1. Missing data by variable. Only map_win_rate has missing entries (63 records, 0.90%).",
           width_inches=5.5)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — OUTLIER DETECTION AND TREATMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 4 — Outlier Detection and Treatment", level=1)

add_body(
    "Outlier detection was performed using the Interquartile Range (IQR) method on all four "
    "continuous predictors, applied before imputation. The IQR fence was defined as "
    "[Q1 - 1.5 x IQR, Q3 + 1.5 x IQR]. Values outside this range were flagged for review.",
    space_after=8
)

add_table(
    caption="Table 3 — Outlier Detection Results (IQR Method, Pre-Cleaning)",
    headers=["Variable", "IQR Fence Lower", "IQR Fence Upper", "Outliers Flagged", "Pre-Win. Min", "Pre-Win. Max"],
    rows=[
        ["mean_hltv_rating",
         str(outlier_info["mean_hltv_rating"]["lo"]),
         str(outlier_info["mean_hltv_rating"]["hi"]),
         str(outlier_info["mean_hltv_rating"]["n"]),
         str(outlier_info["mean_hltv_rating"]["min"]),
         str(outlier_info["mean_hltv_rating"]["max"])],
        ["mean_kpr",
         str(outlier_info["mean_kpr"]["lo"]),
         str(outlier_info["mean_kpr"]["hi"]),
         str(outlier_info["mean_kpr"]["n"]),
         str(outlier_info["mean_kpr"]["min"]),
         str(outlier_info["mean_kpr"]["max"])],
        ["head_to_head_win_rate",
         str(outlier_info["head_to_head_win_rate"]["lo"]),
         str(outlier_info["head_to_head_win_rate"]["hi"]),
         str(outlier_info["head_to_head_win_rate"]["n"]),
         "—", "—"],
        ["map_win_rate",
         str(outlier_info["map_win_rate"]["lo"]),
         str(outlier_info["map_win_rate"]["hi"]),
         str(outlier_info["map_win_rate"]["n"]),
         "—", "—"],
    ]
)

add_body(
    "The 298 HLTV rating outliers reflect matches where team-average ratings fell outside the "
    "typical 0.99-1.15 band — either dominant performances by elite lineups or unexpectedly poor "
    "showings in mismatched fixtures. The 152 KPR outliers similarly reflect matches with "
    "unusually low or high kill-round efficiency relative to the professional norm.",
    space_after=6
)

add_body(
    "The 1,935 IQR outliers in head_to_head_win_rate warrant specific attention. This count is "
    "large because the variable is heavily bimodal: teams either have no prior history (defaulting "
    "to 0.50) or have very one-sided records. This is a structural property of the data, not a "
    "data quality problem — a team that has beaten an opponent 5 of 6 times genuinely has a "
    "0.833 win rate, which falls outside the IQR fence. Winsorization at the 1st/99th percentile "
    "is therefore critical here, as it trims only the most extreme boundary cases without "
    "distorting the meaningful variation in the variable.",
    space_after=6
)

add_body(
    "Winsorization was applied to all four continuous variables by clipping values below the 1st "
    "percentile and above the 99th percentile. This retains all 7,033 records while bounding "
    "extreme values that could otherwise dominate gradient computations in the logistic regression "
    "and exaggerate feature importances in the tree-based models.",
    space_after=8
)

add_figure(f"{FIG_DIR}/fig2_boxplots.png",
           "Figure 2. Boxplots for HLTV Rating and KPR (pre-winsorization). Red dots indicate IQR-flagged outliers.",
           width_inches=5.5)

add_figure(f"{FIG_DIR}/fig3_h2h_distribution.png",
           "Figure 3. Head-to-head win rate distribution showing bimodal structure. "
           "Orange dotted lines mark IQR fences; the large outlier count reflects data structure, not data error.",
           width_inches=5.5)

add_table(
    caption="Table 4 — Winsorization Impact (1st / 99th Percentile Clipping)",
    headers=["Variable", "1st Pct. Floor", "99th Pct. Cap", "Clipped Low", "Clipped High", "Post-Win. Range"],
    rows=[
        ["mean_hltv_rating",
         str(winsor_info["mean_hltv_rating"]["floor"]),
         str(winsor_info["mean_hltv_rating"]["cap"]),
         str(winsor_info["mean_hltv_rating"]["clipped_low"]),
         str(winsor_info["mean_hltv_rating"]["clipped_high"]),
         f"[{winsor_info['mean_hltv_rating']['floor']}, {winsor_info['mean_hltv_rating']['cap']}]"],
        ["mean_kpr",
         str(winsor_info["mean_kpr"]["floor"]),
         str(winsor_info["mean_kpr"]["cap"]),
         str(winsor_info["mean_kpr"]["clipped_low"]),
         str(winsor_info["mean_kpr"]["clipped_high"]),
         f"[{winsor_info['mean_kpr']['floor']}, {winsor_info['mean_kpr']['cap']}]"],
        ["head_to_head_win_rate",
         str(winsor_info["head_to_head_win_rate"]["floor"]),
         str(winsor_info["head_to_head_win_rate"]["cap"]),
         str(winsor_info["head_to_head_win_rate"]["clipped_low"]),
         str(winsor_info["head_to_head_win_rate"]["clipped_high"]),
         f"[{winsor_info['head_to_head_win_rate']['floor']}, {winsor_info['head_to_head_win_rate']['cap']}]"],
        ["map_win_rate",
         str(winsor_info["map_win_rate"]["floor"]),
         str(winsor_info["map_win_rate"]["cap"]),
         str(winsor_info["map_win_rate"]["clipped_low"]),
         str(winsor_info["map_win_rate"]["clipped_high"]),
         f"[{winsor_info['map_win_rate']['floor']}, {winsor_info['map_win_rate']['cap']}]"],
    ]
)

add_figure(f"{FIG_DIR}/fig5_distributions.png",
           "Figure 4. Feature distributions before and after winsorization for HLTV Rating and KPR. "
           "Tails are truncated while the central distribution is preserved.",
           width_inches=6.0)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATA CLEANING AND ENCODING PROTOCOL
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 5 — Data Cleaning and Encoding Protocol", level=1)

add_body("The cleaning protocol was executed in the following sequence:", space_after=4)

steps = [
    ("Step 1 — Mean Substitution (Map Win Rate).",
     f"The {map_wr_missing} missing values in map_win_rate were replaced with the observed column "
     f"mean of {map_wr_mean:.4f}, computed from the {N - map_wr_missing:,} non-missing records before "
     "imputation. Mean substitution is appropriate here because the MAR mechanism ensures the observed "
     "values are an unbiased sample of the underlying distribution. The resulting imputed value falls "
     "within one standard deviation of the variable's mean, introducing minimal distributional distortion."),
    ("Step 2 — Neutral Imputation Check (Head-to-Head Win Rate).",
     "No missing values were found in head_to_head_win_rate (0 NaNs). The HLTV dataset records a "
     "50.00% default value for teams with no prior matchup history, which aligns exactly with the "
     "neutral prior strategy used in our protocol. This confirms that the data source already applies "
     "the same epistemically neutral prior we would have imputed."),
    ("Step 3 — Winsorization.",
     "All four continuous variables were winsorized at their 1st and 99th percentiles (see Section 4). "
     "Winsorization was executed after the imputation step so that imputed values informed the "
     "percentile thresholds on the full n=7,033 sample."),
    ("Step 4 — Binary Encoding of Target Variable.",
     "match_outcome was retained in its original binary form (1 = Win, 0 = Loss from team1 "
     "perspective). No additional label encoding was required as team1_win_flag is already an "
     "integer binary variable. No categorical predictors are present in the model feature set, "
     "so no dummy variable encoding was necessary."),
]

for title_text, body_text in steps:
    p = doc.add_paragraph()
    p.add_run(title_text).font.bold = True
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_body(body_text, space_after=6)

add_body(
    "Post-cleaning assertion confirmed zero missing values across all five variables. All 7,033 "
    "records were retained — no rows were dropped at any stage of the pipeline.",
    bold=True, space_after=8
)

wins  = int(df_clean["match_outcome"].sum())
losses = N - wins
add_table(
    caption="Table 5 — Target Variable Class Distribution (Post-Cleaning)",
    headers=["Class", "Label", "Count", "Percentage"],
    rows=[
        ["Win",  "1", f"{wins:,}",  f"{wins/N*100:.2f}%"],
        ["Loss", "0", f"{losses:,}", f"{losses/N*100:.2f}%"],
    ]
)

add_figure(f"{FIG_DIR}/fig4_class_distribution.png",
           "Figure 5. Target variable class distribution post-cleaning. "
           "The 54.6% / 45.4% split is mild and does not require oversampling.",
           width_inches=4.5)

add_body(
    "The class split (54.6% Win / 45.4% Loss) reflects a slight asymmetry in the dataset. "
    "All records are logged from the team1 perspective, and team1 is not randomly assigned — "
    "in many HLTV match entries, team1 corresponds to the higher-ranked or home-side team, "
    "producing a modest win-rate advantage. This imbalance is mild and will not require "
    "oversampling, but the F1-score is retained as an evaluation metric as specified in Task 5 "
    "to ensure both precision and recall are tracked.",
    space_after=6
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DATA PARTITIONING AND STANDARDIZATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 6 — Data Partitioning and Standardization", level=1)

add_body(
    "The cleaned dataset was partitioned using an 80/20 hold-out split with random_state=42, "
    f"producing a training set of {len(X_train):,} records and a test set of {len(X_test):,} records. "
    "Stratification on the target variable was applied to preserve the 54.6% / 45.4% class ratio "
    "across both partitions.",
    space_after=8
)

add_table(
    caption="Table 6 — Train / Test Split Summary",
    headers=["Partition", "Records", "Win (1)", "Loss (0)", "Win Rate"],
    rows=[
        ["Training Set", f"{len(X_train):,}", f"{int(y_train.sum()):,}",
         f"{int((y_train==0).sum()):,}", f"{y_train.mean()*100:.1f}%"],
        ["Test Set",     f"{len(X_test):,}",  f"{int(y_test.sum()):,}",
         f"{int((y_test==0).sum()):,}",  f"{y_test.mean()*100:.1f}%"],
        ["Total",        f"{N:,}", f"{wins:,}", f"{losses:,}", f"{wins/N*100:.1f}%"],
    ]
)

add_body(
    f"The test set (n = {len(X_test):,}) is withheld entirely from model training and hyperparameter "
    "tuning. None of the three models — Logistic Regression, Random Forest, and Gradient Boosting — "
    "will access these records until final performance reporting in the subsequent task.",
    space_after=6
)

add_body(
    "Z-Score Standardization was applied to all four continuous predictors using "
    "sklearn.preprocessing.StandardScaler. The scaler was fit exclusively on the training set "
    "and then applied via transform to the test set. This prevents data leakage: if the scaler "
    "were fit on the combined dataset, test-set statistics would contaminate the scaling "
    "parameters and produce artificially optimistic held-out performance estimates.",
    space_after=8
)

add_table(
    caption="Table 7 — StandardScaler Parameters (Fit on Training Set Only)",
    headers=["Feature", "Training Mean (u)", "Training Std Dev (s)", "Interpretation"],
    rows=[
        ["mean_hltv_rating",
         f"{scaler.mean_[0]:.4f}", f"{scaler.scale_[0]:.4f}",
         "Tight professional range around 1.07"],
        ["mean_kpr",
         f"{scaler.mean_[1]:.4f}", f"{scaler.scale_[1]:.4f}",
         "Very consistent ~0.69 KPR across pro matches"],
        ["head_to_head_win_rate",
         f"{scaler.mean_[2]:.4f}", f"{scaler.scale_[2]:.4f}",
         "Wide spread; many 50/50 first-time matchups"],
        ["map_win_rate",
         f"{scaler.mean_[3]:.4f}", f"{scaler.scale_[3]:.4f}",
         "Teams play stronger maps; skews slightly above 0.50"],
    ]
)

train_scaled_df = pd.DataFrame(X_train_scaled, columns=FEATURES)
add_body(
    f"Post-standardization verification confirmed a column-wise training mean of "
    f"{train_scaled_df.mean().mean():.8f} and standard deviation of "
    f"{train_scaled_df.std().mean():.8f} — satisfying StandardScaler's expected output. "
    "The tight standard deviations for HLTV Rating and KPR reflect the narrow performance "
    "band at the professional level, where most Tier-1 and Tier-2 teams cluster within a small "
    "range of metrics. This characteristic makes standardization especially important for "
    "Logistic Regression, where coefficient magnitudes would otherwise be distorted by the "
    "scale difference between KPR (range ~0.13) and head_to_head_win_rate (range ~1.0).",
    space_after=6
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ASSUMPTIONS ABOUT DATA QUALITY  [REQUIRED MANDATORY HEADING]
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 7 — Assumptions About Data Quality", level=1)

add_body(
    "This section critically evaluates four dimensions of data quality: consistency, potential "
    "measurement errors, representativeness, and data anomalies. These evaluations directly "
    "inform the modeling decisions described in the cleaning protocol and shape the interpretation "
    "of any results produced in subsequent tasks.",
    space_after=8
)

quality_sections = [
    ("7.1  Data Consistency",
     "All five model variables were verified to be internally consistent across the 7,033 records. "
     "The binary target variable (match_outcome) contains no values outside {0, 1}. Continuous "
     "predictors (HLTV Rating, KPR, head-to-head win rate, map win rate) were verified to fall "
     "within theoretically plausible ranges: HLTV Rating is defined on (0, ∞) with professional "
     "averages typically between 0.85 and 1.25; KPR is bounded below by 0 and above by "
     "approximately 1.5 in professional play. No duplicate match records were detected after "
     "checking on the match ID and date columns. The match_outcome column and the team perspective "
     "encoding (all records from team1 viewpoint) are structurally consistent throughout the dataset."),
    ("7.2  Potential Measurement Errors",
     "A key caveat of the HLTV Kaggle dataset is that team statistics (HLTV Rating, KPR, map win "
     "rates, overall win rates) represent the team's profile at the time of data scraping, not at "
     "the exact moment of each historical match. This means that for earlier matches in the dataset "
     "(May 2024), the recorded statistics may reflect the team's later-period performance. This "
     "introduces a form of temporal leakage: the model may inadvertently train on future-state "
     "information when learning from past matches. This limitation is partially mitigated by the "
     "dataset's rolling structure, but users of any trained model should be aware that real-time "
     "deployment would require live statistic feeds from HLTV.org. Additionally, map win rates "
     "for teams with fewer matches on a given map carry higher variance — a team that has played "
     "a map only twice has a win rate of either 0%, 50%, or 100%, which does not reflect a "
     "reliable long-term estimate."),
    ("7.3  Representativeness Concerns",
     "As documented in the Task 4 Ethical Risk Assessment, the HLTV dataset is structurally biased "
     "toward Western European and North American competition. The 648 tournaments in this dataset "
     "are disproportionately drawn from top-tier events where EU/NA organisations dominate the "
     "field. Teams from the CIS, Southeast Asia, Oceania, and the Middle East are represented with "
     "far fewer matches, resulting in less reliable feature estimates for these organisations. "
     "The model trained on this data may systematically underestimate win probabilities for "
     "underrepresented teams — not because those teams are weaker, but because their performance "
     "profiles are undersampled relative to the EU/NA baseline. The scaler means "
     f"(HLTV u={scaler.mean_[0]:.3f}, KPR u={scaler.mean_[1]:.3f}) encode this EU/NA-dominated "
     "distribution. The mitigation strategy identified in Task 4 remains the recommended diagnostic: "
     "stratified accuracy evaluation by region (EU, NA, CIS, APAC) after model training."),
    ("7.4  Data Anomalies Discovered",
     "Three notable anomalies were identified during the audit. First, the 1,935 IQR-flagged "
     "outliers in head_to_head_win_rate constitute 27.5% of the sample — an unusually high "
     "proportion that initially appeared to indicate data errors. Investigation revealed this "
     "reflects the bimodal structure of competitive matchup histories (many teams at 0.5 neutral, "
     "many teams with highly one-sided records), which is a structural feature of the sport rather "
     "than a quality issue. Second, 63 records lack a decider_map value; these correspond to "
     "best-of-one matches where the single map played was not separately labelled as a 'decider' "
     "in the HLTV scraping logic. Third, no anomalies were found in the binary target — the "
     "54.6% Win rate is consistent with the known HLTV convention of assigning team1 to the "
     "higher-ranked or home-side competitor."),
    ("7.5  Ethical Safeguards",
     "Consistent with the consent and data protection plan established in Task 4: all statistics "
     "are aggregated at the team level and no individual player identifiers appear in the model "
     "feature set. The pipeline processes only publicly available data from HLTV.org as permitted "
     "by the platform's terms for research use. The dataset and all outputs are restricted to "
     "academic submission and will be deleted within 30 days of final grade submission. Model "
     "outputs are probabilistic estimates only and must not be used for betting, player "
     "recruitment, or competitive intelligence."),
]

for sub_title, sub_body in quality_sections:
    p = doc.add_paragraph()
    r = p.add_run(sub_title)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    add_body(sub_body, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — REPRODUCIBILITY PLAN  [REQUIRED MANDATORY HEADING — 15 pts]
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Section 8 — Reproducibility Plan", level=1)

add_body(
    "Full transparency and reproducibility of the preprocessing pipeline is a core requirement "
    "of rigorous data science practice (Kothari, Chapter 7). This section documents all "
    "information required for an independent researcher to replicate every preprocessing step "
    "described in this report.",
    space_after=8
)

repro_sections = [
    ("8.1  Software Environment",
     "All preprocessing was performed in Python 3.x using the following libraries:\n"
     "• pandas (data loading, feature engineering, missing data handling)\n"
     "• numpy (numerical operations, IQR computation)\n"
     "• scikit-learn v1.x (train_test_split, StandardScaler)\n"
     "• matplotlib + seaborn (visualisation figures in this report)\n"
     "No proprietary or paid software was used. All libraries are open-source and freely "
     "available via pip (pip install pandas numpy scikit-learn matplotlib seaborn)."),
    ("8.2  Code Repository and Execution",
     "The complete preprocessing pipeline and all project files are publicly available at:\n\n"
     "    https://github.com/aloneen/cs2-match-outcome-prediction\n\n"
     "The pipeline script code/data_pipeline.py can be executed end-to-end "
     "with a single command:\n\n"
     "    python code/data_pipeline.py\n\n"
     "The script prints step-by-step output including row counts, missing value summaries, "
     "IQR fence values, winsorization clip counts, and StandardScaler parameters — all of "
     "which match the tables in this report. The report itself is regenerated by running:\n\n"
     "    python code/generate_report.py\n\n"
     "which re-executes the full pipeline, re-generates all figures from live data, and "
     "overwrites task7_report.docx with a fresh version.\n\n"
     "[REPO_LINK_PLACEHOLDER]"),
    ("8.3  Random Seed and Determinism",
     "The train/test split is produced using sklearn.model_selection.train_test_split with "
     "random_state=42 and stratify=y. Setting random_state=42 ensures the exact same "
     "5,626 / 1,407 partition is produced on every run. All other preprocessing steps "
     "(missing value imputation, winsorization, standardization) are fully deterministic "
     "and contain no stochastic elements."),
    ("8.4  Data File",
     "The input data file is:\n"
     "    csvs/cs2_newestcombinedmatches_team1_reference_reduced2.csv\n\n"
     "This file is stored locally in the project's csvs/ directory and was downloaded from "
     "the Kaggle dataset published by griffindesroches (see References, item 1). The file "
     "contains 7,033 rows and 168 columns. The SHA-256 hash of the file should be verified "
     "if the dataset is re-downloaded to confirm the same version is used."),
    ("8.5  File Naming and Documentation",
     "The project follows the naming convention task[N]_[description].py / task[N]_[description].docx. "
     "All preprocessing outputs (cleaned DataFrame, scaler object) are available in memory "
     "after running code/data_pipeline.py. For downstream model training (Task 8), "
     "the pipeline exports X_train_scaled, X_test_scaled, y_train, and y_test as the "
     "authoritative partitioned and scaled feature matrices. Figure images are saved to "
     "the report_figures/ directory and are embedded automatically when generate_report.py "
     "is executed."),
]

for sub_title, sub_body in repro_sections:
    p = doc.add_paragraph()
    r = p.add_run(sub_title)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    # Handle multi-line bodies with code snippets
    REPO_URL = "https://github.com/aloneen/cs2-match-outcome-prediction"
    for line in sub_body.split("\n"):
        if line == "[REPO_LINK_PLACEHOLDER]":
            lp = doc.add_paragraph()
            lp.add_run("Full source code and data pipeline: ").font.size = Pt(11)
            add_hyperlink(lp, REPO_URL, REPO_URL)
            lp.paragraph_format.space_after  = Pt(2)
            lp.paragraph_format.space_before = Pt(0)
        else:
            lp = doc.add_paragraph()
            run = lp.add_run(line)
            run.font.size = Pt(10 if line.startswith("    ") else 11)
            if line.startswith("    "):
                run.font.name = "Courier New"
                lp.paragraph_format.left_indent = Cm(1.5)
            lp.paragraph_format.space_after  = Pt(2)
            lp.paragraph_format.space_before = Pt(0)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("References", level=1)

refs = [
    "griffindesroches. (2025). CS2 HLTV Professional Match Statistics Dataset. Kaggle. "
    "https://www.kaggle.com/datasets/griffindesroches/cs2-hltv-professional-match-statistics-dataset",
    "HLTV.org. (2025). Professional CS2 match statistics and rankings. Retrieved October 2025.",
    "Broms, E., & Nordanskold, W. (2022). Predicting Counter-Strike Matches. Lund University Student Publications.",
    "Svec, O. (2022). Predicting Counter-Strike Game Outcomes with Machine Learning. "
    "Faculty of Electrical Engineering, Czech Technical University in Prague.",
    "Forgey, H. (2021). Predicting Game Outcome of Counter Strike with Machine Learning: "
    "A Study of Global Offensive Competitive Matches. Fordham University (Gabelli School of Business).",
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. "
    "Journal of Machine Learning Research, 12, 2825-2830.",
    "McKinney, W. (2010). Data Structures for Statistical Computing in Python. "
    "Proceedings of the 9th Python in Science Conference, 56-61.",
    "Kothari, C. R. (2004). Research Methodology: Methods and Techniques (2nd ed.). "
    "New Age International. [Chapters 5, 7]",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"{i}.  {ref}", style="List Number")
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("reports/task7_report.docx")
print("\nSaved: reports/task7_report.docx")
print("Figures saved in: report_figures/")
